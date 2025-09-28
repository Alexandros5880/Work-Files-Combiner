import os
import re
from copy import deepcopy
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

from docx import Document

# Προσπάθεια για πλήρη συγχώνευση (εικόνες/formatting) με docxcompose αν υπάρχει
DOCXCOMPOSE_AVAILABLE = False
try:
    from docxcompose.composer import Composer
    DOCXCOMPOSE_AVAILABLE = True
except Exception:
    DOCXCOMPOSE_AVAILABLE = False


def natural_key(s: str):
    """
    Φυσική ταξινόμηση: διασπά το string σε κομμάτια αριθμών/κειμένου για σωστή σειρά
    π.χ. ['01', '2', '10'] -> 1,2,10 (όχι 1,10,2)
    """
    return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', s)]


def numeric_prefix_key(s: str):
    """
    Ταξινόμηση με βάση αριθμητικό πρόθεμα στην αρχή του ονόματος (π.χ. '01_Intro.docx' -> 1).
    Αν δεν βρεθεί, πέφτει πίσω σε φυσική ταξινόμηση.
    """
    m = re.match(r'^\s*(\d+)', s)
    if m:
        return (0, int(m.group(1)), s.lower())
    return (1, ) + tuple(natural_key(s))


def combine_with_docxcompose(files, out_path):
    if not files:
        raise ValueError("Δεν επιλέχθηκαν αρχεία.")
    base = Document(files[0])
    composer = Composer(base)
    for f in files[1:]:
        composer.append(Document(f))
    composer.save(out_path)


def append_heading(doc: Document, text: str):
    h = doc.add_heading(text, level=1)
    # Προαιρετικά: κενό μετά τον τίτλο
    doc.add_paragraph()


def append_doc_text_only(dst: Document, src_path: str):
    """
    Απλή αντιγραφή κειμένου (paragraphs). Δεν εγγυάται εικόνες/headers/footers/πίνακες.
    """
    src = Document(src_path)
    for p in src.paragraphs:
        new_p = dst.add_paragraph()
        # αντιγραφή runs για bold/italic όπου είναι εφικτό
        for r in p.runs:
            nr = new_p.add_run(r.text)
            nr.bold = r.bold
            nr.italic = r.italic
            nr.underline = r.underline


def combine_text_only(files, out_path, add_headings=False, insert_breaks=True):
    if not files:
        raise ValueError("Δεν επιλέχθηκαν αρχεία.")
    out = Document()
    for i, f in enumerate(files):
        if add_headings:
            append_heading(out, os.path.basename(f))
        append_doc_text_only(out, f)
        if insert_breaks and i < len(files) - 1:
            out.add_page_break()
    out.save(out_path)


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Combine Word (.docx) Files")
        self.geometry("700x470")
        self.resizable(False, False)

        self.selected_files = []
        self.selected_folder = None
        self.sort_mode = tk.StringVar(value="numeric_prefix")
        self.add_headings = tk.BooleanVar(value=True)
        self.insert_breaks = tk.BooleanVar(value=True)
        self.use_docxcompose = tk.BooleanVar(value=False)
        self.output_path = tk.StringVar(value="")

        self.create_widgets()

    def create_widgets(self):
        pad = 10

        # Επιλογές Εισόδου
        frm_in = ttk.LabelFrame(self, text="Είσοδος")
        frm_in.place(x=pad, y=pad, width=680, height=160)

        ttk.Button(frm_in, text="Επιλογή .docx αρχείων…", command=self.pick_files).place(x=15, y=15)
        ttk.Button(frm_in, text="Επιλογή φακέλου…", command=self.pick_folder).place(x=200, y=15)
        ttk.Button(frm_in, text="Καθαρισμός λίστας", command=self.clear_selection).place(x=320, y=15)

        self.files_list = tk.Listbox(frm_in, height=5)
        self.files_list.place(x=15, y=55, width=650, height=85)

        # Επιλογές Ταξινόμησης
        frm_sort = ttk.LabelFrame(self, text="Ταξινόμηση")
        frm_sort.place(x=pad, y=pad+170, width=330, height=110)

        ttk.Radiobutton(frm_sort, text="Αριθμητικό πρόθεμα (01, 02…)", value="numeric_prefix",
                        variable=self.sort_mode).place(x=10, y=10)
        ttk.Radiobutton(frm_sort, text="Φυσική αλφαριθμητική (A1, A2… A10)", value="natural",
                        variable=self.sort_mode).place(x=10, y=40)

        # Επιλογές Συγχώνευσης
        frm_opts = ttk.LabelFrame(self, text="Επιλογές συγχώνευσης")
        frm_opts.place(x=360, y=pad+170, width=330, height=110)

        ttk.Checkbutton(frm_opts, text="Heading με το όνομα αρχείου", variable=self.add_headings).place(x=10, y=10)
        ttk.Checkbutton(frm_opts, text="Page break ανάμεσα στα αρχεία", variable=self.insert_breaks).place(x=10, y=40)
        ttk.Checkbutton(frm_opts, text="Πλήρης συγχώνευση (docxcompose)", variable=self.use_docxcompose,
                        state=("normal" if DOCXCOMPOSE_AVAILABLE else "disabled")).place(x=10, y=70)

        # Έξοδος
        frm_out = ttk.LabelFrame(self, text="Έξοδος")
        frm_out.place(x=pad, y=pad+290, width=680, height=90)

        ttk.Button(frm_out, text="Αποθήκευση ως…", command=self.pick_output).place(x=15, y=15)
        ttk.Entry(frm_out, textvariable=self.output_path).place(x=140, y=17, width=530)

        # Κουμπιά Εκτέλεσης
        ttk.Button(self, text="ΣΥΓΧΩΝΕΥΣΗ", command=self.run_merge).place(x=470, y=pad+390, width=200, height=40)

        # Κατάσταση
        self.status = tk.StringVar(value=("Έτοιμο. " +
                         ("(docxcompose διαθέσιμο)" if DOCXCOMPOSE_AVAILABLE else "(docxcompose ΜΗ διαθέσιμο)")))
        ttk.Label(self, textvariable=self.status, anchor="w").place(x=pad, y=pad+390, width=430, height=40)

    def pick_files(self):
        files = filedialog.askopenfilenames(title="Επίλεξε .docx αρχεία",
                                            filetypes=[("Word Documents", "*.docx")])
        if files:
            self.selected_files = list(files)
            self.selected_folder = None  # ακυρώνει επιλογή φακέλου
            self.refresh_list()

    def pick_folder(self):
        folder = filedialog.askdirectory(title="Επίλεξε φάκελο με .docx")
        if folder:
            self.selected_folder = folder
            # συλλογή όλων των .docx στο φάκελο (όχι υποφακέλους)
            self.selected_files = [os.path.join(folder, f) for f in os.listdir(folder) if f.lower().endswith(".docx")]
            self.refresh_list()

    def clear_selection(self):
        self.selected_files = []
        self.selected_folder = None
        self.refresh_list()

    def refresh_list(self):
        self.files_list.delete(0, tk.END)
        # ταξινόμηση
        if self.sort_mode.get() == "numeric_prefix":
            sorted_files = sorted(self.selected_files, key=lambda p: numeric_prefix_key(os.path.basename(p)))
        else:
            sorted_files = sorted(self.selected_files, key=lambda p: natural_key(os.path.basename(p)))

        self.selected_files = sorted_files
        for f in self.selected_files:
            self.files_list.insert(tk.END, os.path.basename(f))

        self.status.set(f"Επιλεγμένα: {len(self.selected_files)}")

    def pick_output(self):
        out = filedialog.asksaveasfilename(
            title="Αποθήκευση ως",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile="combined.docx"
        )
        if out:
            self.output_path.set(out)

    def run_merge(self):
        try:
            if not self.selected_files:
                messagebox.showwarning("Προσοχή", "Δεν έχεις επιλέξει αρχεία ή φάκελο.")
                return
            if not self.output_path.get():
                self.pick_output()
                if not self.output_path.get():
                    return

            out_path = self.output_path.get()

            if self.use_docxcompose.get() and DOCXCOMPOSE_AVAILABLE:
                self.status.set("Συγχώνευση (docxcompose)…")
                self.update_idletasks()
                combine_with_docxcompose(self.selected_files, out_path)
            else:
                self.status.set("Συγχώνευση (κείμενο μόνο)…")
                self.update_idletasks()
                combine_text_only(
                    self.selected_files,
                    out_path,
                    add_headings=self.add_headings.get(),
                    insert_breaks=self.insert_breaks.get()
                )

            self.status.set(f"Ολοκληρώθηκε → {out_path}")
            messagebox.showinfo("Έτοιμο", f"Δημιουργήθηκε το αρχείο:\n{out_path}")

        except Exception as ex:
            self.status.set("Σφάλμα")
            messagebox.showerror("Σφάλμα", f"Κάτι πήγε στραβά:\n{ex}")


if __name__ == "__main__":
    App().mainloop()
